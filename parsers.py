import re, unicodedata, zipfile
from typing import Dict, Tuple

# ---------------- Exceptions ----------------
class PDFSupportMissing(Exception):
    """Se lanza cuando se intenta leer PDF sin pdfplumber instalado."""
    pass

# ---------------- Normalización -------------
def _normalize(s: str) -> str:
    s = s.replace("\xa0", " ")  # NBSP
    s = unicodedata.normalize("NFD", s).encode("ascii", "ignore").decode("ascii")
    s = re.sub(r"[ \t]+", " ", s.lower())
    s = re.sub(r"\r?\n\s*\r?\n+", "\n", s).strip()
    return s

# ---------------- Lectores -------------------
def extract_text_from_docx(path: str) -> str:
    """DOCX robusto: docx2txt (mejor con tablas) -> python-docx (párrafos+tablas) -> XML crudo."""
    # 1) docx2txt
    try:
        import docx2txt
        txt = docx2txt.process(path)
        if txt and txt.strip():
            return txt
    except Exception:
        pass
    # 2) python-docx (incluye tablas)
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if t:
                        parts.append(t)
        if parts:
            return "\n".join(parts)
    except Exception:
        pass
    # 3) XML crudo
    try:
        with zipfile.ZipFile(path) as z:
            xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
        xml = re.sub(r"</w:p>", "\n", xml)
        text = re.sub(r"<[^>]+>", "", xml)
        return text
    except Exception:
        return ""

def extract_text_from_pdf(path: str) -> str:
    try:
        import pdfplumber  # lazy import; opcional
    except Exception as e:
        raise PDFSupportMissing("pdfplumber no esta instalado.") from e
    texts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            texts.append(page.extract_text() or "")
    return "\n".join(texts)

def extract_text(file_path: str) -> Tuple[str, str]:
    """Devuelve (texto_normalizado, formato) para .docx/.pdf/.txt"""
    lower = (file_path or "").lower()
    if lower.endswith(".docx"):
        raw = extract_text_from_docx(file_path)
        return _normalize(raw), "docx"
    elif lower.endswith(".pdf"):
        raw = extract_text_from_pdf(file_path)
        return _normalize(raw), "pdf"
    elif lower.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return _normalize(f.read()), "txt"
    elif lower.endswith(".doc"):
        raise ValueError("Formato .DOC (Word 97-2003) no soportado. Convertir a .DOCX o PDF.")
    else:
        raise ValueError("Formato no soportado. Use .docx, .pdf o .txt")

# ---------------- Helpers -------------------
def _count(rx: str, t: str) -> int:
    return len(re.findall(rx, t, flags=re.IGNORECASE))

def _has(rx: str, t: str) -> int:
    return 1 if re.search(rx, t, flags=re.IGNORECASE) else 0

def _num(rx: str, t: str) -> int:
    m = re.search(rx, t, flags=re.IGNORECASE)
    return int(m.group(1)) if m else 0

# ---------------- Detectores -----------------
def detect_counts(text: str) -> Dict[str, int]:
    t = _normalize(text)
    c: Dict[str, int] = {}

    # ---------- FORMACION ----------
    c["formacion:doctorado"] = _has(r"\bdoctorado\b", t)
    c["formacion:maestria"] = _count(r"\bmaestria\b", t)
    c["formacion:especializacion"] = _count(r"\bespecializacion\b", t)
    c["formacion:diplomatura"] = _count(r"\bdiplomatura\b", t)
    c["formacion:segundo_grado"] = _has(r"segundo titulo de grado|doble titulo de grado", t)
    c["formacion:cursos_posgrado"] = _num(r"(?:cursos? de (?:pos|post)grado|cursos? de especializacion).*?(\d+)", t)
    c["formacion:posdoc"] = _count(r"\bposdoc(?:torado)?\b", t)
    c["formacion:idiomas"] = _count(r"\bidioma[s]?\b|ingles certificado|toefl|ielts|b2|c1|c2", t)
    c["formacion:estancias"] = _num(r"(?:estancia|pasantia)[^0-9]{0,20}(\d+)", t) or _count(r"\bestancia\b|\bpasantia\b", t)

    # ---------- DOCENCIA ----------
    def _years(rx):
        m = re.search(rx, t)
        return int(m.group(1)) if m else 0
    c["docencia:titular"] = _years(r"titular[^0-9]{0,20}(\d+)\s*a(?:n|ñ)os?")
    c["docencia:asociado"] = _years(r"asociad[oa][^0-9]{0,20}(\d+)\s*a(?:n|ñ)os?")
    c["docencia:adjunto"] = _years(r"adjunt[oa][^0-9]{0,20}(\d+)\s*a(?:n|ñ)os?")
    c["docencia:jtp"] = _years(r"(?:trabajos practicos|ayudante)[^0-9]{0,20}(\d+)\s*a(?:n|ñ)os?")
    c["docencia:posgrado"] = _years(r"posgrado[^0-9]{0,20}(\d+)\s*cursos?")

    # ---------- GESTION ----------
    c["gestion:rector"] = _has(r"\brector(?:a|ado)?\b", t)
    c["gestion:vicerrector"] = _has(r"\bvicerrector(?:a)?\b|directorio", t)
    c["gestion:decano"] = _has(r"\bdecano\b|director[ae] de facultad|director instituto", t)
    c["gestion:secretario"] = _has(r"secretari[oa] (academica|de investigacion|de extension)", t)
    c["gestion:coordinador"] = _has(r"coordinador[ae] de carrera|responsable de programas", t)
    c["gestion:consejero"] = _has(r"consejer[oa] (superior|directivo|de facultad)", t)

    # ---------- OTROS CARGOS ----------
    c["otroscargos:funciones"] = _num(r"comisiones? internas?[^0-9]{0,20}(\d+)\s*funciones?", t) or _count(r"comision interna", t)

    # ---------- FORMACION RRHH ----------
    c["ciencia:dir_doctorandos"] = _num(r"direccion de (\d+) doctorand", t)
    c["ciencia:dir_maestria"] = _num(r"direccion de (\d+) maestrand", t)
    c["ciencia:dir_grado"] = _num(r"direccion de (\d+) tesis(?:tas)? de grado", t)
    c["ciencia:becarios"] = _num(r"(\d+)\s*becarios? (?:conicet|agencia)", t)

    # ---------- PROYECTOS ----------
    c["proyectos:direccion"] = _num(r"(?:direccion|dir\.) de (\d+) proyectos?", t)
    c["proyectos:codireccion"] = _num(r"co-?direccion de (\d+) proyectos?", t)
    c["proyectos:participacion"] = _num(r"participacion en (\d+) proyectos?", t)
    c["proyectos:coordinacion"] = _num(r"coordinacion de (\d+) equipo", t)

    # ---------- EXTENSION ----------
    c["extension:tutorias"] = _num(r"tutorias? de (\d+) pasant", t)
    c["extension:transferencia"] = _num(r"(\d+)\s*(?:actividades?|acciones?) de transferencia", t)
    c["extension:eventos_cientificos"] = _num(r"(\d+)\s*(?:conferencias?|panel(?:es)?|exposiciones?)", t)

    # ---------- EVALUACION ----------
    c["eval:tribunal_grado"] = _num(r"jurado de (\d+) tesis de grado", t)
    c["eval:tribunal_posgrado"] = _num(r"jurado de (\d+) tesis de posgrado", t)
    c["eval:eval_revistas"] = max(
        _num(r"evaluador(?:a)? de (\d+) (?:revistas?|congresos?|jornadas?)", t),
        _count(r"\b(revisor|arbitro|reviewer)\b", t)
    )
    c["eval:eval_proyectos"] = max(
        _num(r"evaluador(?:a)? de (\d+) proyectos?\s*(?:i\+d|investigacion)?", t),
        _count(r"\bevaluador(?:a)? de proyectos\b", t)
    )
    c["eval:eval_institucional"] = max(
        _num(r"evaluacion institucional.*?(\d+)", t),
        _count(r"\bevaluacion institucional\b", t)
    )

    # ---------- OTRAS ACTIVIDADES ----------
    c["otras:comites_redes"] = max(
        _num(r"participacion en (\d+) redes", t),
        _count(r"red(?:es)? academicas?|comite(s)?", t)
    )
    c["otras:ejercicio_prof"] = _num(r"extraacademico.*?(\d+)\s*a(?:n|ñ)os?", t)

    # ---------- PUBLICACIONES (PRODUCCIONES) ----------
    # Heurísticas robustas para CVs: DOI/ISSN/ISBN por línea, sin exigir la palabra "articulo".
    lines = t.split("\n")

    # DOIs únicos (muy confiables para artículos)
    dois = set(re.findall(r"\b10\.\d{4,9}/[-._;()/:a-z0-9]+\b", t, flags=re.I))

    # ISBN y ISSN marcados
    # (permitimos espacios entre letras en PDFs raros: i s b n / i s s n)
    isbn_token = re.compile(r"i\s*s\s*b\s*n", re.I)
    issn_token = re.compile(r"i\s*s\s*s\s*n", re.I)

    libro_count = 0
    cap_count = 0
    issn_lines = 0

    for ln in lines:
        has_isbn = bool(isbn_token.search(ln))
        has_issn = bool(issn_token.search(ln))
        if has_isbn:
            if re.search(r"\bcapitul|chapter\b", ln, flags=re.I):
                cap_count += 1
            else:
                libro_count += 1
        if has_issn:
            # Evito confundir con líneas de libro que traigan ISSN incidental
            if not has_isbn and not re.search(r"\bcapitul|chapter|libro\b", ln, flags=re.I):
                issn_lines += 1

    # Artículos con referato: preferimos DOIs, si no ISSN por línea,
    # y como respaldo, líneas con 'journal|revista' sin ISBN.
    journal_lines = 0
    for ln in lines:
        if re.search(r"\b(journal|revista)\b", ln, flags=re.I) and not isbn_token.search(ln):
            journal_lines += 1

    c["pubs:con_referato"] = max(len(dois), issn_lines, journal_lines)

    # Artículos sin referato: palabras clave típicas
    c["pubs:sin_referato"] = max(
        _num(r"(\d+)\s*articulos?\s*sin\s*referato", t),
        _count(r"(divulgacion|boletin|articulo de opinion|nota periodistica)", t)
    )

    # Libros y capítulos por ISBN en línea
    c["pubs:libros"] = libro_count
    c["pubs:capitulos"] = cap_count

    # Documentos técnicos / informes
    c["pubs:documentos"] = max(
        _num(r"(\d+)\s*(?:documentos?|informes?)\s*tecnicos?", t),
        _count(r"(informe|documento|manual|guia).{0,20}tecnic", t)
    )

    # ---------- DESARROLLOS ----------
    c["desarrollos:software_patente"] = max(
        _num(r"(\d+)\s*softwares?\s*registrados?", t),
        _num(r"(\d+)\s*patentes?", t)
    )
    c["desarrollos:procesos"] = _num(r"(\d+)\s*procesos?\s*de\s*gestion", t)

    # ---------- SERVICIOS ----------
    c["servicios:tecnicos"] = _num(r"(\d+)\s*servicios?\s*tecnicos?", t)
    c["servicios:informes"] = _num(r"(\d+)\s*informes?\s*tecnicos?", t)

    # ---------- REDES / EDITORIAL / EVENTOS ----------
    c["redes:participacion"] = max(
        _num(r"miembro de (\d+) redes?", t),
        _count(r"\b(membresia|miembro|socio|asociado|integrante)\b", t)
    )
    c["redes:organizacion_eventos"] = max(
        _num(r"organizacion de (\d+) (?:congresos?|jornadas?|seminarios?)", t),
        _count(r"(comite (?:organizador|cientifico)|coordinacion) de (?:congreso|jornada|seminario|evento)", t)
    )
    c["redes:gestion_editorial"] = max(
        _num(r"comite editorial de (\d+) revistas?", t),
        _count(r"\b(editor(?:a)?(?: asociado[a]?| en jefe)?|comite editorial)\b", t)
    )

    # ---------- PREMIOS ----------
    total_premios = _count(r"\bpremio[s]?\b|\bdistincion(?:es)?\b|\bmencion(?:es)?\b|\breconocimiento[s]?\b", t)
    premios_int = _count(r"\binternac", t)
    premios_nac = _count(r"\bnacional", t)
    c["premios:internacional"] = premios_int
    c["premios:nacional"] = max(premios_nac - premios_int, 0)
    resto = max(total_premios - c["premios:internacional"] - c["premios:nacional"], 0)
    c["premios:distinciones"] = resto

    return c
