from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _p(p, text, bold=False, size=11):
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def build_docx_report(df_items, totals, meta=None) -> bytes:
    """
    df_items: DataFrame con columnas: Sección, Ítem, Unidades detectadas,
              Puntos por unidad, Puntos (tope ítem)
    totals: dict con claves como '2:Formacion_total', ..., 'TOTAL_GENERAL'
    meta: dict opcional: {'docente': str, 'institucion': str}
    """
    meta = meta or {}
    doc = Document()

    # Portada
    title = doc.add_heading('Informe de Valoración – Categorizador Docente en Investigación', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    _p(p, f"Fecha: {datetime.now().strftime('%d/%m/%Y')}", size=10)
    if meta.get("docente"):
        p = doc.add_paragraph()
        _p(p, f"Docente: {meta['docente']}", bold=True)
    if meta.get("institucion"):
        p = doc.add_paragraph()
        _p(p, f"Institución: {meta['institucion']}")

    # Totales por bloque
    doc.add_paragraph()
    doc.add_heading("Totales por bloque", level=1)
    orden = ["2:Formacion_total","3:Cargos_total","4:CyT_total","5:Producciones_total","6:Otros_total","TOTAL_GENERAL"]
    etiquetas = {
        "2:Formacion_total":"Formación",
        "3:Cargos_total":"Cargos (Docencia/Gestión)",
        "4:CyT_total":"Ciencia y Tecnología",
        "5:Producciones_total":"Producciones",
        "6:Otros_total":"Otros",
        "TOTAL_GENERAL":"TOTAL GENERAL"
    }
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text = "Bloque"
    hdr[1].text = "Puntaje"
    for k in orden:
        row = t.add_row().cells
        row[0].text = etiquetas.get(k, k)
        row[1].text = f"{float(totals.get(k,0)):.0f}"

    # Desglose por ítem
    doc.add_paragraph()
    doc.add_heading("Desglose por ítem", level=1)
    cols = ["Sección","Ítem","Unidades detectadas","Puntos por unidad","Puntos (tope ítem)"]
    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Table Grid"
    for i,c in enumerate(cols):
        tbl.rows[0].cells[i].text = c

    for _, r in df_items.iterrows():
        row = tbl.add_row().cells
        row[0].text = str(r.get("Sección",""))
        row[1].text = str(r.get("Ítem",""))
        row[2].text = str(r.get("Unidades detectadas",""))
        row[3].text = str(r.get("Puntos por unidad",""))
        row[4].text = str(r.get("Puntos (tope ítem)",""))

    # Observaciones
    doc.add_paragraph()
    doc.add_heading("Observaciones", level=1)
    p = doc.add_paragraph()
    _p(p, "Este informe fue generado automáticamente a partir del CV cargado en la app. "
           "Las detecciones se basan en expresiones regulares y pueden requerir validación humana.")

    # Exportar a bytes
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
